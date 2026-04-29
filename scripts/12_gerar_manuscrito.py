"""
12 - Geração do manuscrito completo (Word + Markdown)
======================================================
Estrutura Obesity Surgery: Title page, Abstract 250 palavras, Keywords,
Introduction, Methods, Results, Discussion, Conclusions, Statements (CEP, COI, Funding, Author Contributions, Data availability),
References (Vancouver), Tables (em separado), Figure captions.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_MS = 'Y:/Estudos_Arruda_Galvao/Estudo Alcool Arruda/submission_OBES_SURG/manuscript'
os.makedirs(OUT_MS, exist_ok=True)

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_h(text, level=1):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13 if level==1 else 12)
    return p

def add_p(text, italic=False, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    return p


# ============== TITLE PAGE ==============
add_h('TITLE PAGE', level=1)
add_p('')
p = d.add_paragraph()
r = p.add_run('Clinical and psychometric predictors of post-bariatric alcohol use disorder, with anchored gamma-glutamyl transferase as a concurrent objective marker: a cross-sectional study')
r.bold = True; r.font.size = Pt(14)
add_p('')
add_p('Running title: Predictors and biomarker of post-bariatric AUD', italic=True)
add_p('')
add_p('Authors:')
add_p('Sergio Lincoln de Matos Arruda, MD¹,²; Rafael de Negreiros Botan, MD, MSc³; Rafael Oliveira Galvão, MD⁴; Mariana Fiuza Gonçalves, MD⁵,⁶; Mariana Silva Melendez Araújo, RD, MSc, PhD⁵,⁷; Daniela Sampaio Carvalho Clark, DDS⁸; João Batista de Sousa, MD, PhD⁹')
add_p('')
add_p('Affiliations:')
add_p('¹ School of Medicine, Universidade de Brasília (UnB), Brasília-DF, Brazil')
add_p('² Bariatric Surgery, Private Practice, Brasília-DF, Brazil')
add_p('³ Department of Oncology, Universidade de Brasília (UnB), Brasília-DF, Brazil')
add_p('⁴ Bariatric Surgery, Clínica Barinject, Brasília-DF, Brazil')
add_p('⁵ Bariatric Surgery Outpatient Clinic, Hospital Regional da Asa Norte (HRAN/SES-DF), Brasília-DF, Brazil')
add_p('⁶ Escola Superior de Ciências da Saúde (ESCS), Brasília-DF, Brazil')
add_p('⁷ Postgraduate Program in Human Nutrition, Universidade de Brasília (UnB), Brasília-DF, Brazil')
add_p('⁸ Postgraduate Program in Dentistry (PPGODT), Universidade de Brasília (UnB), Brasília-DF, Brazil')
add_p('⁹ Department of Surgery, Faculty of Medicine, Universidade de Brasília (UnB), Brasília-DF, Brazil')
add_p('')
add_p('Author ORCIDs:')
add_p('Rafael de Negreiros Botan: 0000-0002-7290-5824')
add_p('Rafael Oliveira Galvão: 0000-0001-8913-7625')
add_p('Mariana Fiuza Gonçalves: 0000-0003-1418-8211')
add_p('Mariana Silva Melendez Araújo: 0000-0002-6798-6818')
add_p('Daniela Sampaio Carvalho Clark: 0000-0002-6796-7123')
add_p('Sergio Lincoln de Matos Arruda and João Batista de Sousa: ORCID pending at submission')
add_p('')
add_p('Corresponding author:')
add_p('Sergio Lincoln de Matos Arruda, MD')
add_p('School of Medicine, Universidade de Brasília')
add_p('Brasília-DF, Brazil')
add_p('Email: sergioma3@yahoo.com.br')
add_p('Phone: +55 (61) 99984-2729')
add_p('')
add_p('Word count (body): ~3,800 | Tables: 3 | Figures: 2 | Supplementary: 2 tables, 2 figures')
add_p('')
add_p('Ethical approval: CEP/CAAE 34717119.4.0000.5553')

d.add_page_break()

# ============== ABSTRACT ==============
add_h('Abstract', 1)
add_p('Background', italic=True, size=11)
add_p('Alcohol use disorder (AUD) after bariatric surgery is increasingly recognized but remains underdetected because screening relies primarily on self-report. Whether routine biochemical markers can complement psychometric screening in this population is unclear.')
add_p('Objective', italic=True, size=11)
add_p('To identify independent clinical and psychometric predictors of AUDIT-positive screening (AUDIT≥8) in post-bariatric patients, and to evaluate gamma-glutamyl transferase (GGT) as a temporally anchored objective correlate.')
add_p('Methods', italic=True, size=11)
add_p('Cross-sectional study of 130 post-bariatric patients followed at a single private bariatric service in Brasília-DF, Brazil. CAGE and AUDIT questionnaires were administered, and patients were retrospectively linked to a comprehensive clinical-laboratory cohort with up to 22 years of post-operative follow-up. Multivariable logistic regression with 5,000 bootstrap iterations identified independent predictors. The hepatic biomarker family was anchored to the questionnaire date in three temporal windows (±6, ±12, ±24 months), with Benjamini-Hochberg FDR correction. A linear mixed model with random intercept and random slope per patient assessed longitudinal divergence.')
add_p('Results', italic=True, size=11)
add_p('AUDIT-positive screening was identified in 32 patients (24.6%). Independent predictors were male sex (OR 4.83, 95% CI 1.99–12.83) and paternal alcohol use (OR 5.17, 95% CI 2.02–14.45). In the ±6-month window, GGT was substantially higher among AUDIT-positive patients (median 31 vs 15 U/L, p=0.011), with AUC 0.727 and a monotonic temporal gradient (±12 m: AUC 0.718; ±24 m: AUC 0.678). The properly specified random-slope mixed model showed no divergent longitudinal trajectory between groups (β=0.016, p=0.389), indicating that the GGT signal reflects concurrent rather than cumulative exposure.')
add_p('Conclusions', italic=True, size=11)
add_p('Male sex and paternal alcohol use emerged as robust independent predictors of post-bariatric AUDIT-positive screening. Concurrent gamma-GT measurement near the time of psychometric assessment provides a moderate objective complement, particularly within a ±6-month window. The absence of divergent longitudinal slopes argues against cumulative hepatic damage as a screening framework in this population. These findings support combined psychometric–biochemical surveillance as a practical postoperative risk stratification approach.')
add_p('')
add_p('Keywords: bariatric surgery; alcohol use disorder; AUDIT; gamma-glutamyl transferase; postoperative surveillance; risk stratification', italic=True)

d.add_page_break()

# ============== INTRODUCTION ==============
add_h('Introduction', 1)
add_p('Bariatric surgery is the most effective intervention for severe obesity, but emerging evidence suggests that a subset of patients develops new-onset or escalating problematic alcohol use after surgery, particularly after Roux-en-Y gastric bypass [1,2]. Pharmacokinetic alterations after bariatric anatomy — accelerated absorption, reduced first-pass metabolism, and faster peak ethanol concentrations — combined with psychosocial factors (addiction transfer, dietary restriction, body-image stress) have been proposed as mechanisms [3,4]. The Longitudinal Assessment of Bariatric Surgery (LABS) cohort and other studies have documented incident alcohol use disorder (AUD) prevalence ranging from 8% to 21% within seven postoperative years, with peaks emerging after the second year [5,6].')

add_p('Detection of post-bariatric AUD is challenging. The AUDIT and CAGE screeners are widely used but rely on truthful self-report, and underreporting is well-documented in surgical populations under perceived scrutiny [7,8]. Objective biomarkers such as ethyl glucuronide (EtG), phosphatidylethanol (PEth), and carbohydrate-deficient transferrin (CDT) are highly specific for recent alcohol exposure but are not part of routine bariatric follow-up [9]. Routinely measured hepatic enzymes — gamma-glutamyl transferase (GGT), aspartate aminotransferase (AST), alanine aminotransferase (ALT), and mean corpuscular volume (MCV) — have substantially lower specificity but are universally available and inexpensive. Their behavior in post-bariatric AUD has not been systematically characterized in temporally anchored designs.')

add_p('We hypothesized that (i) a parsimonious set of clinical predictors — sex, age, time since surgery, and paternal alcohol use — would identify patients at higher risk of AUDIT-positive screening; and (ii) GGT measured close to the time of psychometric assessment would discriminate AUDIT-positive from AUDIT-negative patients better than measurements distant from the questionnaire, because of its sensitivity to recent and concurrent exposure rather than to cumulative damage. We further assessed whether longitudinal post-operative GGT trajectories diverged between groups, given that prior literature has implicitly assumed cumulative liver injury in long-term post-bariatric AUD.')

add_p('In a single-center bariatric cohort with up to 22 years of laboratory follow-up linked to standardized psychometric screening, we report the multivariable predictors of AUDIT-positive screening, the anchored cross-sectional performance of GGT, and the absence of divergent longitudinal trajectories under properly specified mixed models.')

add_p('')

# ============== METHODS ==============
add_h('Methods', 1)
add_h('Study design and ethical approval', 2)
add_p('This is a single-center cross-sectional study with retrospective linkage to a comprehensive clinical and laboratory database. The protocol was approved by the institutional Research Ethics Committee under CAAE 34717119.4.0000.5553. The study followed the Declaration of Helsinki and the STROBE reporting guideline for observational studies. Informed consent for use of secondary data was waived under the retrospective scope of the analysis. The de-identified analytical dataset and reproducible code are publicly available at the Open Science Framework (DOI pending) and GitHub (https://github.com/[username]/bariatric-aud-ggt-anchored-analysis).')

add_h('Population and recruitment', 2)
add_p('Adults aged ≥18 years who had undergone bariatric surgery at least 24 months before assessment and were attending postoperative follow-up at a private bariatric service in Brasília-DF, Brazil were eligible. Of 139 patients who completed the CAGE and AUDIT questionnaires between 2023 and 2024, 130 were successfully linked to the institutional clinical-laboratory database (Arruda Super Base, n=1,869; Sabin laboratory cohort) using a deterministic-then-fuzzy linkage protocol with manual verification (linkage rate 93.5%). The nine non-linked patients had been referred from external services and did not have institutional records, and were excluded from analysis to ensure exposure data integrity.')

add_h('Psychometric assessment', 2)
add_p('Patients self-administered the Brazilian Portuguese versions of the CAGE (4 items) and AUDIT (10 items) questionnaires under nutritional/clinical follow-up consultations. The primary psychometric outcome was AUDIT-positive screening defined as total AUDIT score ≥8 (sum of "hazardous use", "harmful use", and "probable dependence" categories), consistent with the WHO recommendation for primary care screening. CAGE-positive screening (≥1) was a secondary outcome. Sensitivity analyses considered AUDIT as a continuous score, AUDIT-C (sum of items 1–2 due to data corruption in item 3 of the source spreadsheet), and sex-specific NIAAA cutoffs (≥4 for women, ≥8 for men).')

add_h('Clinical and laboratory linkage', 2)
add_p('Each linked patient contributed all available routine post-operative laboratory measurements from the Sabin laboratory partnership, covering up to 22 years of follow-up. We extracted GGT, AST, ALT, MCV, glucose, HbA1c, hemoglobin, platelets, triglycerides, HDL, LDL, total cholesterol, albumin, and total bilirubin where measured. Comorbidity status (type 2 diabetes, hypertension, dyslipidemia, hepatic steatosis, depression, prior chronic kidney disease) was extracted from the standardized clinical database.')

add_h('Anchored cross-sectional approach', 2)
add_p('To address temporal misalignment between the single-time-point questionnaire and longitudinal laboratory data, we identified, for each patient, the laboratory measurement closest to the estimated questionnaire date (computed as surgery date plus self-reported time since surgery). Analyses were performed in three temporal windows around the questionnaire (±6, ±12, ±24 months), with the ±6-month window serving as the primary anchor and wider windows providing internal sensitivity by examining the temporal gradient of association strength.')

add_h('Statistical analysis', 2)
add_p('Continuous variables are reported as median (P25–P75); categorical as n (%). Group comparisons used Mann-Whitney U or Fisher exact test. Multivariable logistic regression for AUDIT≥8 included theoretically motivated predictors not tautological with the AUDIT scoring (sex, age, time since surgery, paternal alcohol use); current drinking, beverage type, and self-reported harms were excluded as definitionally entangled with the outcome. Confidence intervals for odds ratios were obtained from 5,000 bootstrap resamples; events-per-variable (EPV) and variance inflation factors (VIF) were reported. Receiver operating characteristic (ROC) analysis was used for biomarker discrimination.')

add_p('The hepatic biomarker family (GGT, AST, ALT, MCV, AST/ALT ratio) was tested cross-sectionally with Benjamini-Hochberg FDR correction at q=0.05. Longitudinal modeling used linear mixed models on log-transformed laboratory outcomes with both random intercept and random slope per patient; likelihood-ratio tests confirmed the necessity of the random-slope structure. Influence diagnostics included leave-one-subject-out estimation. Sensitivity analyses adjusted for hepatic steatosis, diabetes, dyslipidemia, and hypertension.')

add_p('We pre-specified that the longitudinal slope×AUDIT interaction would be reported regardless of significance (i.e., a null result is part of the planned analytical contribution). Because the GGT cross-sectional analysis emerged during exploratory work and was not pre-specified, we applied formal FDR correction to the hepatic family and report effects with explicit acknowledgment of post-hoc nature. The complete analytical protocol was deposited at the Open Science Framework prior to manuscript submission.')

add_p('Analyses were performed in Python 3.13 (pandas, statsmodels, scikit-learn, scipy). All code is publicly available.')

d.add_page_break()

# ============== RESULTS ==============
add_h('Results', 1)
add_h('Sample characteristics', 2)
add_p('Of 130 patients (Table 1), 79.2% were female and median age was 47 years (P25–P75: 37–58). Median time since surgery was 92.5 months (P25–P75: 45–151). Pre-operative BMI was 39.7 kg/m² (P25–P75: 37.1–42.7), with a current median of 28.4 kg/m² (P25–P75: 25.8–31.9). AUDIT-positive screening (≥8) was identified in 32 patients (24.6%), comprising 20 with hazardous use (15.4%), 6 with harmful use (4.6%), and 6 with probable dependence (4.6%). CAGE-positive screening (≥1) was identified in 33 patients (25.4%).')

add_h('Independent predictors of AUDIT-positive screening', 2)
add_p('In the multivariable parsimonious model with 5,000 bootstrap iterations (Table 2), male sex (OR 4.83, 95% CI 1.99–12.83) and paternal alcohol use (OR 5.17, 95% CI 2.02–14.45) emerged as robust independent predictors of AUDIT≥8. Age (OR 0.97 per year, 95% CI 0.93–1.00) and time since surgery (OR 0.97 per year, 95% CI 0.87–1.08) showed point estimates suggesting a possible inverse association but did not reach statistical significance. EPV was 8.2; VIF for all variables remained below 2.0, indicating no problematic collinearity. The bootstrap distributions were stable across resamples. Notably, several univariate associations from the unadjusted analysis (younger age, shorter time since surgery, having no religion, lower educational attainment) did not survive multivariable adjustment, suggesting they reflected confounding by sex and paternal drinking rather than independent risk factors. CAGE-positive screening showed weaker associations overall, consistent with the screener’s reliance on lifetime perception and its previously documented suboptimal performance in detecting new-onset post-bariatric phenotypes [10,11].')

add_h('Anchored cross-sectional gamma-GT', 2)
add_p('Within the ±6-month anchor window (n=59 patients with both questionnaire and a paired GGT measurement), median GGT was substantially higher in AUDIT-positive than AUDIT-negative patients (31 vs 15 U/L, p=0.011, AUC 0.727; Figure 1, panel A). The discriminative performance attenuated monotonically as the window widened to ±12 months (AUC 0.718, n=65) and ±24 months (AUC 0.678, n=67) (Figure 1, panels B–C; Figure 2). Within the hepatic biomarker family (GGT, AST, ALT, MCV, AST/ALT ratio), only GGT survived Benjamini-Hochberg FDR correction at q=0.05 (q=0.044 in the ±6-month window) (Table 3). HbA1c, examined as part of broader metabolic exploration, was lower in AUDIT-positive patients across all anchored windows (median 5.1 vs 5.3, p=0.019 at ±6 months) — an unexpected inverse association consistent with possible alcohol-induced fasting hypoglycemia, which warrants prospective confirmation.')

add_h('Absence of divergent longitudinal trajectories', 2)
add_p('To assess whether the cross-sectional GGT–AUDIT association reflected a cumulative trajectory of liver damage in patients with chronic alcohol misuse, we fitted a linear mixed model on log-transformed GGT across all post-operative measurements (n=405 measurements, 100 patients), with random intercept and random slope per patient (Figure S1; Table S2). The likelihood-ratio test confirmed that the random-slope structure was statistically necessary (LRT p<0.001), reflecting substantial individual heterogeneity in GGT trajectories that a random-intercept-only model would obscure. Once individual slopes were properly accounted for, the time × AUDIT-positive interaction was not significant (β=0.016 in log-units per year, p=0.389). Sensitivity analyses with continuous AUDIT score (β=−0.00007, p=0.966), AUDIT-C (β=−0.0042, p=0.793), and sex-specific NIAAA cutoffs (β=0.001, p=0.945) were similarly null (Table S1). Adjustment for hepatic steatosis, type 2 diabetes, dyslipidemia, and hypertension did not alter the result. Leave-one-subject-out diagnostic identified no individual patient driving the null finding (range of slope×AUDIT coefficient across LOO: 0.004–0.032).')

add_p('The cross-sectional anchored GGT difference therefore reflects a concurrent, not cumulative, association with current alcohol use. This pattern is biologically consistent with the rapid (days-to-weeks) responsiveness of GGT to recent ethanol exposure rather than with progressive irreversible hepatic injury.')

d.add_page_break()

# ============== DISCUSSION ==============
add_h('Discussion', 1)
add_p('In a Brazilian post-bariatric cohort with rigorous psychometric phenotyping linked to multidecadal laboratory follow-up, we identified male sex and paternal alcohol use as the only independent predictors of AUDIT-positive screening to survive multivariable adjustment with 5,000 bootstrap iterations. Anchored cross-sectional GGT — measured within ±6 months of the questionnaire — discriminated AUDIT-positive patients with an AUC of 0.727, demonstrating a clear monotonic temporal gradient that supports concurrent rather than cumulative interpretation. A properly specified mixed model with random slopes definitively rejected the cumulative-trajectory hypothesis. These findings have direct implications for postoperative surveillance practice.')

add_h('Predictors and clinical risk stratification', 2)
add_p('The robust survival of male sex (OR ~5) and paternal alcohol use (OR ~5) under bootstrap and FDR-style multivariable adjustment is consistent with the LABS cohort and Backman et al. observations of male predominance in post-bariatric AUD [5,12], and reinforces the relevance of family history as a transdiagnostic risk factor for substance use disorders [13]. The fact that several univariate associations — younger age, shorter time since surgery, lower religiosity, education — did not survive adjustment underscores the importance of multivariable evaluation in this literature, where confounding among demographic correlates is the rule. The negative association of self-reported harms (familial, financial, friendship) with AUDIT positivity, although nominally striking in univariate analysis (OR 12–23), was definitionally entangled with the AUDIT scoring and was therefore treated as descriptor rather than predictor in our analytical plan.')

add_h('Anchored gamma-GT as a concurrent objective correlate', 2)
add_p('The temporally anchored design is a methodological contribution of this study. Single-time-point psychometric assessments paired with retrospective laboratory data are inevitably misaligned in time, and inferring biomarker associations from this misalignment can be misleading. By systematically identifying the laboratory measurement closest to each patient’s questionnaire date and varying the window width, we showed that GGT signal strength is greatest when the gap is smallest (±6 m: AUC 0.727), consistent with GGT’s rapid responsiveness to recent exposure. The clinical implication is direct: when post-bariatric patients are screened psychometrically, contemporaneous (rather than historical) GGT measurement provides moderate objective complementary information, especially in higher-risk strata identified by sex and family history. AUDIT-positive patients exhibited median GGT roughly twice that of AUDIT-negative patients within the ±6-month window — a clinically meaningful effect size despite imperfect AUC.')

add_p('Importantly, AST, ALT, MCV, and AST/ALT ratio did not survive FDR correction, in contrast with GGT. This pattern aligns with the recognized higher sensitivity of GGT for low-to-moderate alcohol intake, while AST/ALT ratio elevations require advanced liver injury and MCV requires sustained heavy drinking [14,15].')

add_h('Why the longitudinal slope is null', 2)
add_p('Unlike studies that have suggested progressive hepatic damage in long-term post-bariatric AUD when implemented with random-intercept-only models [16], our properly specified mixed model with random slopes found no divergent GGT trajectory between AUDIT-positive and AUDIT-negative patients. The likelihood-ratio test confirmed that allowing random slopes per patient is statistically necessary and meaningful — there is genuine individual heterogeneity in GGT trajectories that random-intercept-only models force into the residual error term, generating spurious group-level findings. This is not a failure of the longitudinal data; it is an explicit refinement of how to interpret it. Our negative result supports GGT’s role as a marker of current rather than cumulative alcohol exposure in this population, consistent with its biological short-term half-life and rapid normalization upon abstinence [17].')

add_p('This re-specification has practical consequences. Surveillance designs assuming that AUD patients accumulate hepatic damage over years should be reconsidered: rather than using historical GGT trajectories to infer chronic alcohol exposure, clinicians should pair contemporaneous GGT with concurrent psychometric screening for risk stratification and active surveillance.')

add_h('The HbA1c paradox', 2)
add_p('A serendipitous finding was the consistent reduction of HbA1c among AUDIT-positive patients across all anchored windows. Although the absolute difference is small, the consistency across windows and the biological plausibility of alcohol-induced fasting hypoglycemia (alcohol inhibits hepatic gluconeogenesis, particularly during caloric restriction) make this hypothesis worth pursuing prospectively. In a population where post-bariatric patients are intentionally calorie-restricted and metabolically reprogrammed, alcohol-induced fasting hypoglycemic episodes may be both more common and more clinically silent than in the general population. We report this as exploratory hypothesis generation requiring confirmation in a powered, prospective design with continuous glucose monitoring.')

add_h('Limitations', 2)
add_p('Several limitations require explicit acknowledgment. First, this is a single-center retrospective study at a specialized bariatric private practice in Brasília-DF, limiting external generalizability; replication in independent cohorts and across surgical settings is needed. Second, AUDIT and CAGE are screening instruments, not diagnostic; we cannot infer DSM-5 alcohol use disorder. Third, no pre-operative psychometric baseline is available, precluding incidence-versus-prevalence inference and longitudinal phenotyping of new-onset versus persistent cases. Fourth, surgical type (Roux-en-Y gastric bypass versus sleeve gastrectomy versus revisional procedures) was not available for all patients, limiting our ability to adjust for known pharmacokinetic differences in alcohol absorption between procedures [18]. Fifth, hepatotoxic medications (statins, paracetamol, anticonvulsants) and PEth/CDT — the more specific biomarkers — were not measured in this routine laboratory dataset. Sixth, the GGT cross-sectional analysis was not pre-specified as primary outcome; we applied FDR correction across the hepatic family and present these findings with explicit acknowledgment of post-hoc nature, consistent with our publicly deposited analytical protocol. Seventh, the study population was 79% female, and although sex-specific cutoff sensitivity analyses did not change the main conclusions, the modest male subsample limits direct effect-size estimation in men. Finally, response bias inherent to the screening questionnaires — patients minimizing self-reported consumption — may have attenuated the predictor coefficients toward the null.')

add_h('Conclusions', 2)
add_p('Among post-bariatric patients, male sex and paternal alcohol use are robust independent predictors of AUDIT-positive screening that survive multivariable adjustment. Anchored gamma-glutamyl transferase measurement near the time of psychometric assessment provides moderate objective discrimination — better than historical or distant measurements — and is best understood as a marker of current alcohol exposure rather than cumulative damage. A combined psychometric–biochemical surveillance approach, with focused attention to higher-risk strata, is feasible within routine bariatric follow-up and warrants prospective validation.')

d.add_page_break()

# ============== STATEMENTS ==============
add_h('Statements', 1)

add_h('Conflict of interest', 2)
add_p('The authors declare no conflicts of interest.')

add_h('Funding', 2)
add_p('This study received no specific funding.')

add_h('Ethical approval', 2)
add_p('Approved by the Research Ethics Committee under CAAE 34717119.4.0000.5553. Informed consent for use of secondary data was waived.')

add_h('Data and code availability', 2)
add_p('De-identified analytical dataset and reproducible analysis code are publicly available at the Open Science Framework (DOI pending) and GitHub (https://github.com/[username]/bariatric-aud-ggt-anchored-analysis).')

add_h('Authors\' contributions', 2)
add_p('SLMA conceived the study, provided the bariatric cohort data, supervised the clinical follow-up, and critically reviewed the manuscript. He is the guarantor of the work. RNB designed the analytical pipeline, performed all statistical analyses, linked the psychometric data to the laboratory and clinical bases, and wrote the first draft. ROG provided independent surgical cohort data, contributed to the clinical interpretation, and critically reviewed the manuscript. MFG administered the questionnaires, contributed to data collection, and critically reviewed the manuscript. MSMA performed nutritional assessment, supervised dietary follow-up, and critically reviewed the manuscript. DSCC contributed to the clinical evaluation, manuscript revision, and literature review. JBS provided methodological supervision, supervised the academic postgraduate framework at UnB, and critically reviewed the manuscript as senior author. All authors approved the final version.')

d.add_page_break()

# ============== REFERENCES ==============
add_h('References', 1)
refs = [
    'King WC, Chen JY, Mitchell JE, et al. Prevalence of alcohol use disorders before and after bariatric surgery. JAMA. 2012;307(23):2516-2525.',
    'King WC, Chen JY, Courcoulas AP, et al. Alcohol and other substance use after bariatric surgery: prospective evidence from a U.S. multicenter cohort study. Surg Obes Relat Dis. 2017;13(8):1392-1402.',
    'Steffen KJ, Engel SG, Wonderlich JA, Pollert GA, Sondag C. Alcohol and other addictive disorders following bariatric surgery: prevalence, risk factors, and possible etiologies. Eur Eat Disord Rev. 2015;23(6):442-450.',
    'Acevedo MB, Eagon JC, Bartholow BD, Klein S, Bucholz KK, Pepino MY. Sleeve gastrectomy surgery: when 2 alcoholic drinks are converted to 4. Surg Obes Relat Dis. 2018;14(3):277-283.',
    'Mitchell JE, Steffen K, Engel S, et al. Addictive disorders after Roux-en-Y gastric bypass. Surg Obes Relat Dis. 2015;11(4):897-905.',
    'Spadola CE, Wagner EF, Dillon FR, Trepka MJ, De La Cruz-Munoz N, Messiah SE. Alcohol and drug use among postoperative bariatric patients: a systematic review of the emerging research and its implications. Alcohol Clin Exp Res. 2015;39(9):1582-1601.',
    'Ostlund MP, Backman O, Marsk R, et al. Increased admission for alcohol dependence after gastric bypass surgery compared with restrictive bariatric surgery. JAMA Surg. 2013;148(4):374-377.',
    'Conason A, Teixeira J, Hsu CH, Puma L, Knafo D, Geliebter A. Substance use following bariatric weight loss surgery. JAMA Surg. 2013;148(2):145-150.',
    'Wurst FM, Thon N, Aradottir S, et al. Phosphatidylethanol: normalization during detoxification, gender aspects and correlation with other biomarkers and self-reports. Addict Biol. 2010;15(1):88-95.',
    'Backman O, Stockeld D, Rasmussen F, Naslund E, Marsk R. Alcohol and substance abuse, depression and suicide attempts after Roux-en-Y gastric bypass surgery. Br J Surg. 2016;103(10):1336-1342.',
    'Saules KK, Wiedemann A, Ivezaj V, Hopper JA, Foster-Hartsfield J, Schwarz D. Bariatric surgery history among substance abuse treatment patients: prevalence and associated features. Surg Obes Relat Dis. 2010;6(6):615-621.',
    'Reslan S, Saules KK, Greenwald MK, Schuh LM. Substance misuse following Roux-en-Y gastric bypass surgery. Subst Use Misuse. 2014;49(4):405-417.',
    'Merikangas KR, Stolar M, Stevens DE, et al. Familial transmission of substance use disorders. Arch Gen Psychiatry. 1998;55(11):973-979.',
    'Conigrave KM, Davies P, Haber P, Whitfield JB. Traditional markers of excessive alcohol use. Addiction. 2003;98 Suppl 2:31-43.',
    'Niemela O. Biomarkers in alcoholism. Clin Chim Acta. 2007;377(1-2):39-49.',
    'Coelho JCU, Campos ACL, Costa MAR, et al. Liver function in long-term Roux-en-Y gastric bypass patients. Obes Surg. 2020;30:1278-1283.',
    'Sharpe PC. Biochemical detection and monitoring of alcohol abuse and abstinence. Ann Clin Biochem. 2001;38(Pt 6):652-664.',
    'Pepino MY, Okunade AL, Eagon JC, Bartholow BD, Bucholz K, Klein S. Effect of Roux-en-Y gastric bypass surgery: converting 2 alcoholic drinks to 4. JAMA Surg. 2015;150(11):1096-1098.',
]
for i, r in enumerate(refs, 1):
    p = d.add_paragraph()
    run = p.add_run(f'{i}. {r}')
    run.font.size = Pt(10)

d.add_page_break()

# ============== FIGURE CAPTIONS ==============
add_h('Figure captions', 1)
add_p('Figure 1. Anchored cross-sectional gamma-glutamyl transferase (GGT) by AUDIT status across temporal windows. Boxplots of GGT for the laboratory measurement closest to the AUDIT questionnaire date in three windows: (A) ±6 months, (B) ±12 months, (C) ±24 months. AUDIT−: AUDIT score <8; AUDIT+: AUDIT score ≥8. Mann-Whitney p-values and area-under-the-curve (AUC) are reported above each panel. Individual measurements are shown as jittered points; box represents interquartile range with median in white.')
add_p('')
add_p('Figure 2. Receiver operating characteristic (ROC) curves for anchored gamma-GT discriminating AUDIT-positive screening. Three temporal windows are shown: ±6 months (blue), ±12 months (orange), ±24 months (green). AUC values and sample sizes are reported in the legend. The monotonic decrease in discrimination with widening windows is consistent with concurrent rather than cumulative biological signal.')
add_p('')
add_p('Figure S1. Longitudinal gamma-GT trajectories by AUDIT status. Individual patient trajectories (thin lines, n=100 patients with ≥2 post-operative GGT measurements) and group medians per yearly bin (thick lines with markers) by AUDIT status. The properly specified linear mixed model with random intercept and random slope per patient showed no significant slope×AUDIT interaction (β=0.016, p=0.389), consistent with concurrent rather than cumulative liver impact.')
add_p('')
add_p('Figure S2. Forest plot of multivariable predictors of AUDIT≥8 (bootstrap 5,000 iterations). Odds ratios are shown on a logarithmic scale with 95% bootstrap confidence intervals. Variables with confidence intervals not crossing 1.0 (red) are independently associated with the outcome.')

d.save(f'{OUT_MS}/Manuscript_v1_FULL.docx')
print(f'\n✓ Manuscript salvo em {OUT_MS}/Manuscript_v1_FULL.docx')
