"""
13 - Cover letter + STROBE checklist
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_CL = 'Y:/Estudos_Arruda_Galvao/Estudo Alcool Arruda/submission_OBES_SURG/cover_letter'
OUT_CK = 'Y:/Estudos_Arruda_Galvao/Estudo Alcool Arruda/submission_OBES_SURG/checklists'
os.makedirs(OUT_CL, exist_ok=True); os.makedirs(OUT_CK, exist_ok=True)

# ===== COVER LETTER =====
d = Document()
s = d.styles['Normal']; s.font.name='Calibri'; s.font.size=Pt(11)

def p(text, bold=False, italic=False, align=None):
    para = d.add_paragraph(); r = para.add_run(text)
    r.bold=bold; r.italic=italic
    if align: para.alignment = align
    return para

p('Sergio Lincoln de Matos Arruda, MD', bold=True)
p('School of Medicine, Universidade de Brasília')
p('Brasília-DF, Brazil')
p('Email: sergioma3@yahoo.com.br | Phone: +55 (61) 99984-2729')
p('')
p('April 29, 2026')
p('')
p('Editor-in-Chief')
p('Obesity Surgery — The Official Journal of IFSO')
p('Springer')
p('')
p('Dear Editor,')
p('')
p('We are pleased to submit our original research article entitled "Clinical and psychometric predictors of post-bariatric alcohol use disorder, with anchored gamma-glutamyl transferase as a concurrent objective marker: a cross-sectional study" for consideration in Obesity Surgery.')
p('')
p('Alcohol use disorder (AUD) after bariatric surgery is increasingly recognized as a clinically significant complication, but routine surveillance still depends largely on self-report. Whether routinely measured hepatic biomarkers can complement psychometric screening in this population remains unsettled. In a Brazilian post-bariatric cohort with rigorous psychometric phenotyping (CAGE and AUDIT) linked to multidecadal laboratory follow-up (up to 22 years), we performed three contributions that we believe make this work a good fit for your journal:')
p('')
p('1. Multivariable identification of robust independent predictors of AUDIT≥8 — male sex (OR 4.83) and paternal alcohol use (OR 5.17) — using bootstrap 5,000-iteration confidence intervals, providing a parsimonious risk-stratification framework usable in routine bariatric follow-up.')
p('')
p('2. A temporally anchored cross-sectional design that systematically pairs each questionnaire response with the laboratory measurement closest in time, and explicitly tests how discrimination performance varies with the temporal gap. Within ±6 months of the AUDIT, GGT discriminated AUDIT-positive patients with AUC 0.727; the monotonic decrease across wider windows (±12 m, ±24 m) supports a concurrent rather than cumulative interpretation.')
p('')
p('3. An explicit refutation of the cumulative-trajectory hypothesis that has been implicit in some prior literature. With a properly specified linear mixed model including random intercept and random slope per patient, the longitudinal slope×AUDIT interaction was not significant (β=0.016, p=0.389), demonstrating that GGT signal in this population reflects current rather than historical exposure. This methodological note is, in itself, a contribution to how postoperative surveillance is designed.')
p('')
p('Together, these findings support a combined psychometric–biochemical surveillance framework focused on contemporaneous measurement and on patients with identifiable risk factors. The manuscript is ~3,800 words, with 3 main tables, 2 main figures, 2 supplementary tables and 2 supplementary figures.')
p('')
p('We confirm that:')
p('• This work is original and has not been published elsewhere, in whole or in part.')
p('• It is not under consideration by any other journal.')
p('• All authors have read and approved the final version, agree to be accountable for the content, and meet ICMJE authorship criteria.')
p('• The work has institutional ethical approval (CEP/CAAE 34717119.4.0000.5553).')
p('• De-identified data and reproducible code are publicly deposited at the Open Science Framework and GitHub.')
p('• There are no conflicts of interest. No external funding was received.')
p('')
p('We would respectfully suggest, for editorial consideration, that this work may be of interest to readers focused on long-term post-bariatric surveillance, on the integration of routine laboratory data into clinical follow-up, and on the methodological challenges of inferring causal trajectories from misaligned cross-sectional questionnaire data.')
p('')
p('We thank you for considering our manuscript and look forward to your editorial decision.')
p('')
p('Yours sincerely,')
p('')
p('')
p('Sergio Lincoln de Matos Arruda, MD', bold=True)
p('On behalf of all co-authors')

d.save(f'{OUT_CL}/Cover_Letter_Obesity_Surgery.docx')
print(f'✓ Cover letter: {OUT_CL}/Cover_Letter_Obesity_Surgery.docx')


# ===== STROBE CHECKLIST =====
d = Document()
s = d.styles['Normal']; s.font.name='Calibri'; s.font.size=Pt(10)

p_ = d.add_paragraph(); r = p_.add_run('STROBE Statement — Checklist for Cross-Sectional Studies')
r.bold=True; r.font.size=Pt(13)

p_ = d.add_paragraph(); r = p_.add_run('Manuscript: "Clinical and psychometric predictors of post-bariatric alcohol use disorder, with anchored gamma-glutamyl transferase as a concurrent objective marker: a cross-sectional study"')
r.italic=True; r.font.size=Pt(10)
d.add_paragraph()

t = d.add_table(rows=1, cols=4); t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
for i,h in enumerate(['Item #','Section','Recommendation','Page/Section in manuscript']):
    hdr[i].text = h
    for p2 in hdr[i].paragraphs:
        for r in p2.runs: r.bold=True

items = [
    ('1', 'Title and abstract', 'Indicate the study\'s design with a commonly used term in the title or the abstract', 'Title page; Abstract'),
    ('1', 'Title and abstract', 'Provide in the abstract an informative and balanced summary of what was done and what was found', 'Abstract'),
    ('2', 'Introduction — Background/rationale', 'Explain the scientific background and rationale for the investigation being reported', 'Introduction §1–2'),
    ('3', 'Introduction — Objectives', 'State specific objectives, including any prespecified hypotheses', 'Introduction §3'),
    ('4', 'Methods — Study design', 'Present key elements of study design early in the paper', 'Methods — Study design'),
    ('5', 'Methods — Setting', 'Describe the setting, locations, and relevant dates, including periods of recruitment, exposure, follow-up, and data collection', 'Methods — Population and recruitment'),
    ('6', 'Methods — Participants', 'Give the eligibility criteria, and the sources and methods of selection of participants', 'Methods — Population'),
    ('7', 'Methods — Variables', 'Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers', 'Methods — Psychometric assessment, Clinical and laboratory linkage'),
    ('8', 'Methods — Data sources/measurement', 'For each variable of interest, give sources of data and details of methods of assessment (measurement)', 'Methods — Clinical and laboratory linkage'),
    ('9', 'Methods — Bias', 'Describe any efforts to address potential sources of bias', 'Methods — Anchored cross-sectional approach; Discussion — Limitations'),
    ('10', 'Methods — Study size', 'Explain how the study size was arrived at', 'Methods — Population (139 enrolled, 130 linked)'),
    ('11', 'Methods — Quantitative variables', 'Explain how quantitative variables were handled in the analyses', 'Methods — Statistical analysis'),
    ('12', 'Methods — Statistical methods', 'Describe all statistical methods, including those used to control for confounding', 'Methods — Statistical analysis'),
    ('12b', '', 'Describe any methods used to examine subgroups and interactions', 'Methods — Statistical analysis (sensitivity)'),
    ('12c', '', 'Explain how missing data were addressed', 'Methods — Statistical analysis (anchored windows handle misalignment)'),
    ('12d', '', 'If applicable, describe analytical methods taking account of sampling strategy', 'NA — convenience sample, single center'),
    ('12e', '', 'Describe any sensitivity analyses', 'Methods — Sensitivity (AUDIT continuous, AUDIT-C, sex-specific NIAAA, comorbidity adjustment)'),
    ('13', 'Results — Participants', 'Report numbers of individuals at each stage of study (eligible, examined, included)', 'Results — Sample characteristics; Methods (139 → 130)'),
    ('14', 'Results — Descriptive data', 'Give characteristics of study participants (e.g., demographic, clinical, social) and information on exposures and potential confounders', 'Table 1; Results §1'),
    ('15', 'Results — Outcome data', 'Report numbers of outcome events or summary measures', 'Results §1 (24.6% AUDIT+); Tables 1–3'),
    ('16', 'Results — Main results', 'Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision', 'Results §2–3; Tables 2, 3'),
    ('16b', '', 'Report category boundaries when continuous variables were categorized', 'Methods (AUDIT≥8); Tables'),
    ('16c', '', 'If relevant, consider translating estimates of relative risk into absolute risk for a meaningful time period', 'Discussion §2'),
    ('17', 'Results — Other analyses', 'Report other analyses done — e.g., analyses of subgroups, interactions, sensitivity', 'Results §4 (longitudinal null); Tables S1, S2'),
    ('18', 'Discussion — Key results', 'Summarise key results with reference to study objectives', 'Discussion §1'),
    ('19', 'Discussion — Limitations', 'Discuss limitations of the study, taking into account sources of potential bias or imprecision', 'Discussion — Limitations'),
    ('20', 'Discussion — Interpretation', 'Give a cautious overall interpretation of results considering objectives, limitations, multiplicity of analyses, results from similar studies, and other relevant evidence', 'Discussion §2–4'),
    ('21', 'Discussion — Generalisability', 'Discuss the generalisability (external validity) of the study results', 'Discussion — Limitations (single-center)'),
    ('22', 'Other information — Funding', 'Give the source of funding and the role of the funders for the present study', 'Statements — Funding'),
]

for it in items:
    row = t.add_row().cells
    for i, val in enumerate(it):
        row[i].text = val

p_ = d.add_paragraph()
r = p_.add_run('Note: Page references are approximate, based on the submitted manuscript layout. The complete analytical pipeline and de-identified dataset are publicly available (OSF DOI pending, GitHub link in manuscript).')
r.italic=True; r.font.size=Pt(9)

d.save(f'{OUT_CK}/STROBE_Checklist_Filled.docx')
print(f'✓ STROBE: {OUT_CK}/STROBE_Checklist_Filled.docx')
